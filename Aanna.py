import requests
from urllib.parse import urlparse
from bs4 import BeautifulSoup
# pip install beautifulsoup4
import os
import shutil
import re
import time
from datetime import datetime, timedelta
import docx
# pip install python-docx
from io import BytesIO

# fetch the terminal size
terminal_size = shutil.get_terminal_size()
# calculating the x and y coordinates at the center of the terminal
center_x = terminal_size.columns // 2
center_y = terminal_size.lines // 2

# created this function to print output at the middle of the terminal
def print_middle(string):
    print('-' * (center_x*2))
    print(' ' * (center_x - (len(string)//2)), end="")
    for i in string:
        print(i, end='', flush=True)
        time.sleep(0.03)
    print()
    print('-' * (center_x*2))

# created this function to take the input at the middle of the terminal
def input_middle():
    print()
    i=input(' ' * (center_x - 15))
    print()
    return i

# introduction
print()
print_middle('Hi, myself Aanna, a web scraper created by Niladri Das.')

# created this function to check the validity of a url
def is_valid_url(url):
    try:
        result = urlparse(url)
        return all([result.scheme, result.netloc])
    except ValueError:
        return False

# fetch and store the url of the base website
print_middle(':enter the url you want to scrap:')
while True:
    url=input_middle()
    if is_valid_url(url)==True:
        print_middle("valid input")
        break
    else:
        print_middle("invalid input")
        print_middle(":enter again:")
        continue

# fetch and store the domain name of the base website
domain_name=''
try:
    parsed_url=urlparse(url)
    domain_name=parsed_url.netloc
except:
    print_middle("***[invalid base website]***")

# created this function to create and return a valid file name from a link
def create_filename_from_link(link):
    if ("https://"+domain_name) in link:
        filename=str(link).replace("https://"+domain_name,"").lower().strip('/')
        filename=re.sub(r"[^a-z|0-9]", "_", filename)
    else:
        filename=str(link).replace("https://","").lower().strip('/')
        filename=re.sub(r"[^a-z|0-9]", "_", filename)
    return filename

# created this function to fetch and return entire text from a docx file
def get_text_from_doc(filename):
    doc=docx.Document(filename)
    fullText=[]
    for para in doc.paragraphs:
        fullText.append(para.text)
    return '\n'.join(fullText)

# fetch and store the path of the current directory
current_directory=os.getcwd()

print_middle(":our current location:")
print_middle(current_directory)

# create a directory to store the documents
documents_directory=re.sub(r"[^a-z|0-9]", "_", str(url).replace("https://","").lower().strip('/'))
working_directory=current_directory+'/'+documents_directory

try:
    if os.path.exists(working_directory):
        print_middle('it is not a new url for me')
        list_of_documents=[]
        for i in os.listdir(working_directory):
            if i.endswith('.docx'):
                list_of_documents.append(i.replace('.docx',''))
# if there is no files till now in the document_directory then the variable list_of_documents is assigned as None, because it indicates that the script runs first time ever
    else:
        print_middle('it is a new url for me')
        list_of_documents=None
        os.mkdir(working_directory)
except:
    print_middle("***[script is not able to create the directory]***")

# change the path from the current directory to the working directory
os.chdir(working_directory)

print_middle(":our working location:")
print_middle(working_directory)

# created this function to fetch the title, summary and content from a link and store it into a docx file, also check the status of the docx file
def check_and_store_content_from_link(link):
    filename=create_filename_from_link(link)
    if filename=='':
        filename=create_filename_from_link(url)
    doc=docx.Document()
    try:
        response=requests.get(link)
        soup=BeautifulSoup(response.content,'html.parser')

        title=soup.title.text
        try:
            summary = soup.find('meta', attrs={'name': 'description'})['content']
        except:
            summary="not present"
        content=str(soup)

        if list_of_documents is None:
            doc.add_heading('TITLE: '+title, 0)
            doc.add_heading('SUMMARY: '+summary,3)
            doc.add_paragraph('CONTENT: '+content)
            doc.save(filename+'.docx')
            print('{:<30} | {:<60} | {:<90}'.format("created", filename+'.docx', link))
        else:
            if filename in list_of_documents:
                list_of_documents.remove(filename)
                if get_text_from_doc(filename+'.docx') in ('TITLE: '+title+'\n'+'SUMMARY: '+summary+'\n'+'CONTENT: '+content):
                    print('{:<30} | {:<60} | {:<90}'.format("nothing updated", filename+'.docx', link))
                else:
                    doc.add_heading('TITLE: '+title, 0)
                    doc.add_heading('SUMMARY: '+summary,3)
                    doc.add_paragraph('CONTENT: '+content)
                    doc.save(filename+'.docx')
                    print('{:<30} | {:<60} | {:<90}'.format("updated", filename+'.docx', link))
            else:
                doc.add_heading('TITLE: '+title, 0)
                doc.add_heading('SUMMARY: '+summary,3)
                doc.add_paragraph('CONTENT: '+content)
                doc.save(filename+'.docx')
                print('{:<30} | {:<60} | {:<90}'.format("newly created", filename+'.docx', link))
    except:
        if response.status_code==200:
            try:
                response=requests.get(link)
                image=BytesIO(response.content)
                if list_of_documents is None:
                    doc.add_picture(image)
                    doc.save(filename+'.docx')
                    print('{:<30} | {:<60} | {:<90}'.format("created", filename+'.docx', link))
                else:
                    if filename in list_of_documents:
                        list_of_documents.remove(filename)
                        print('{:<30} | {:<60} | {:<90}'.format("nothing updated", filename+'.docx', link))
                    else:
                        doc.add_picture(image)
                        doc.save(filename+'.docx')
                        print('{:<30} | {:<60} | {:<90}'.format("newly created", filename+'.docx', link))
            except:
                print("[error]")
                print(f"permission denied. not able to scrap the page: {link}.")
                print()
        else:
            print("[error]")
            print(f"invalid url. not able to scrap the page: {link}.")
            print()

# create a dictionary to store webpages connected with the base website as key and number of occurence of the respective webpage in the base website as value
pages={url: 1}

# created this function to add key to pages dictionary or update its frequency if the key is already there, also create docx file by calling function check_and_store_content_from_link for each new key in the page dictionary
def add_page(link):
    if link not in pages:
        pages[link]=1
        check_and_store_content_from_link(link)
    else:
        pages[link]+=1

# created this function to fetch all pages from the base website, also add the link and update the frequency of the link into page dictionary by calling function add_page
def extract_and_operate_all_pages(url):
    print_middle('PROCESS STARTED')
    print('{:<30} | {:<60} | {:<90}'.format("STATUS", "FILE NAME", "LINK"))
    try:
        response = requests.get(url)
        soup = BeautifulSoup(response.content, 'html.parser')
        all_anchor=soup.find_all('a')
    except:
        print_middle("***[script is not able to fetch the base website]***")
        print_middle('PROCESS INTERRUPTED')
        all_anchor=[]
    
    for link in all_anchor:
        href = link.get("href")
        if href is None:
            continue
        elif "https://" in href:
            add_page(href)
        else:
            if href.startswith('#') or href.startswith('/#') or href.startswith('//#'):
                add_page(url)
                continue
            elif href.endswith('#'):
                href=href[:-1]
            elif href.startswith("//"):
                href="https:"+href
            elif href.startswith('/'):
                href="https://"+domain_name+href
            else:
                href="https://"+domain_name+'/'+href
            add_page(href)
    
    if list_of_documents!=None:
        for file in list_of_documents:
            os.remove(file+'.docx')
            print('{:<30} | {:<60} | {:<90}'.format("deleted", file+'.docx', 'no link connected'),flush=True)
            time.sleep(3)
    print_middle('PROCESS ENDED')

# this variable is used to maintain the flow of the code later
control=0

# trigger now
while True:
    print_middle(":do you want to trigger it now:")
    print_middle(":enter [yes] or [no]:")
    flag=input_middle()
    try:
        if 'yes' in flag.lower():
            print_middle(":press [Ctrl+C] to interrupt the flow of the code:")
            extract_and_operate_all_pages(url)
            control=1
            if control==1:
                break
        elif 'no' in flag.lower():
            break
        else:
            print_middle("***[invalid entry]***")
            continue
    except:
        print_middle("***[invalid entry]***")
        continue

# trigger weekly
while True:
    print_middle(":do you want to trigger it weekly:")
    print_middle(":enter [yes] or [no]:")
    flag=input_middle()
    try:
        if 'yes' in flag.lower():
            try:
                print_middle("[0 => Monday] [1 => Tuesday] [2 => Wednesday] [3 => Thursday] [4 => Friday] [5 => Saturday] [6 => Sunday]")
                print_middle(":enter weekday [0-6]:")
                weekday=int(input_middle())
                print_middle(":enter hour [0-23]:")
                hour=int(input_middle())
                print_middle(":enter minute [0-59]:")
                minute=int(input_middle())
                if (weekday<0 and weekday>6) or (hour<0 and hour>23) or (minute<0 and minute>59):
                    raise ValueError
                while True:
                    today=datetime.now()
                    counter = ((weekday - today.weekday()) % 7)
                    next_weekday = today + timedelta(days=counter)

                    year=next_weekday.year
                    month=next_weekday.month
                    day=next_weekday.day

                    next_day=datetime(year, month, day, hour, minute, 0)
                    
                    interval=(next_day-today).total_seconds()

                    if interval<0:
                        next_day=next_day+timedelta(days=7)
                        interval=(next_day-today).total_seconds()

                    print_middle(":next trigger will be occured on:")
                    print_middle(str(next_day))
                    print_middle(":press [Ctrl+C] to interrupt the flow of the code:")
                    
                    time.sleep(interval)

                    if control==1:
                        pages={url: 1}
                        list_of_documents=[]
                        for i in os.listdir(working_directory):
                            if i.endswith('.docx'):
                                list_of_documents.append(i.replace('.docx',''))
                    
                    extract_and_operate_all_pages(url)
                    control=1                    
            except:
                print_middle("***[invalid entry]***")
                continue
        elif 'no' in flag.lower():
            break
        else:
            print_middle("***[invalid entry]***")
            continue
    except:
        print_middle("***[invalid entry]***")
        continue

# change the path from the working directory to the current directory
os.chdir(current_directory)

# prints the location of the directory where all the documents are present
if len(os.listdir(working_directory))!=0:
    print_middle(":all the documents are located at:")
    print_middle(working_directory)

else:
    print_middle("***[no documents are created]***")
    shutil.rmtree(working_directory)